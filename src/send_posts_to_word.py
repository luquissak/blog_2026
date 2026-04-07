import sys
import os
from pathlib import Path
from dotenv import load_dotenv
from googleapiclient import sample_tools
import docx
from docx import Document
import html2text
from datetime import datetime

load_dotenv()
BLOG_ID = os.getenv('BLOG_ID')
iso_string = datetime.now().isoformat()
custom_string = datetime.now().strftime('%Y-%m-%d') 
print(f"ISO Format: {iso_string}")
print(f"Folder Date Format: {custom_string}")

# Agora a pasta será "docao2026-04-07", o que é válido e organizado
OUTPUT_DIR = Path(f"docao_{custom_string}")
OUTPUT_DIR.mkdir(exist_ok=True)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.text.run.Run(docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink


def process_post(doc, post, post_day):
    """Adiciona o conteúdo do post ao documento especificado."""
    doc.add_heading(f"{post['title']} - {post_day}", level=1)
    p = doc.add_paragraph("")
    add_hyperlink(p, "Link original: " + post["url"], post["url"])

    text_content = html2text.html2text(post["content"])
    doc.add_paragraph(text_content + "\n")
    doc.add_page_break()  # Opcional: cada post em uma página nova?


def main(argv):
    service, _ = sample_tools.init(
        argv, "blogger", "v3", __doc__, __file__,
        scope="https://www.googleapis.com/auth/blogger"
    )

    blogs_list = service.blogs().listByUser(userId="self").execute()
    docs = {}  # Dicionário para guardar os documentos de cada ano
    total_processed = 0

    for blog in blogs_list.get("items", []):
        print(f"Lendo posts de: {blog['name']}")
        request = service.posts().list(blogId=blog["id"])

        while request is not None:
            posts_doc = request.execute()
            for post in posts_doc.get("items", []):
                # Parsing da data de forma moderna (ISO 8601)
                dt = datetime.fromisoformat(
                    post["published"].replace('Z', '+00:00'))
                post_day = dt.strftime('%d/%m/%Y')
                year = dt.strftime('%Y')

                # Cria o documento do ano se ele ainda não existir no dicionário
                if year not in docs:
                    docs[year] = Document()
                    print(f"--- Iniciando arquivo para o ano {year} ---")

                process_post(docs[year], post, post_day)
                total_processed += 1
                print(f"Processado: {post['title']} ({year})")

            request = service.posts().list_next(request, posts_doc)

    # Salvando os arquivos (Uma única vez por ano!)
    print("\nFinalizando e salvando arquivos...")
    for year, doc in docs.items():
        filename = OUTPUT_DIR / f"Reflexoes_{year}_posts_content.docx"
        doc.save(str(filename))
        print(f"Salvo: {filename}")

    print(
        f"\nSucesso! {total_processed} posts organizados em {len(docs)} arquivos.")


if __name__ == "__main__":
    main(sys.argv)
